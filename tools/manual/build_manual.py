"""Build the two volumes of the SIMF Control Panel operations manual.

    python tools/manual/build_manual.py

Reads
    docs/manuals/source/book.json         the authored content, EN and AR paired
    docs/manuals/source/page-model.json   the facts extracted from the source
    docs/screenshots/manual/*.png         the captured screens

Writes
    docs/manuals/SIMF-CP-Operations-Manual-EN.docx
    docs/manuals/SIMF-CP-Operations-Manual-AR.docx

The two volumes come out of ONE content file with an "en" and an "ar" value on
every block, so they cannot drift apart in structure - a section added to one is
a section added to both, and a missing translation is a build error rather than
a silently English paragraph in the Arabic book.

A referenced screenshot that is not on disk FAILS THE BUILD. The manual this one
replaces carried two image references whose files had never existed, and nothing
noticed for months.
"""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx_kit import (  # noqa: E402
    ARABIC_FONT, LATIN_FONT, MONO_FONT, add_field, mark_dirty_fields,
    paragraph_rtl, run_font, run_rtl, section_rtl, style_rtl, table_rtl,
)

REPO = Path(__file__).resolve().parents[2]
SOURCE = REPO / "docs/manuals/source"
SHOTS = REPO / "docs/screenshots/manual"
OUT = REPO / "docs/manuals"

# How a page is reached when no named permission gates it.
ACCESS_TEXT = {
    "anonymous": {"en": "Anyone, signed in or not",
                  "ar": "أي شخص، سواء سجّل الدخول أم لا"},
    "authenticated": {"en": "Any signed-in user",
                      "ar": "أي مستخدم مسجّل الدخول"},
}

# Screens the Arabic volume is allowed to show in English. It is EMPTY, and that
# is the point: the sign-in screens render once per account and were the last
# English images in the Arabic book, so the documentation database was recreated
# and they were captured again with the interface switched to Arabic BEFORE
# signing in. Any missing Arabic capture is now a fault, not a fallback.
ENGLISH_ONLY = set()

ACCENT = RGBColor(0x1B, 0x3A, 0x5C)
MUTED = RGBColor(0x55, 0x61, 0x70)

# A route, a path or a URL is Latin text. Dropping one into an Arabic paragraph
# lets the bidirectional algorithm move its leading slash to the other end, so
# "/admin/visitors/new" renders as "admin/visitors/new/". The fix is to give that
# text its OWN RUN with right-to-left cleared - not a Unicode directional
# isolate, which Word renders as a visible placeholder glyph instead of obeying.

# What must never be reordered: routes, source paths, Windows paths, host names,
# URLs, configuration variable names and permission codes.
TECHNICAL_PREFIXES = ("/", "src/", "docs/", "tools/", "SIMF_", "ASPNETCORE_",
                      "http://", "https://")
TECHNICAL_SUFFIXES = (".razor", ".cs", ".ps1", ".json")
TECHNICAL_PATTERN = re.compile(
    r"^(?:[A-Za-z]:\\"                                   # D:\System\...
    r"|[A-Za-z0-9_-]+(?:\.[A-Za-z0-9_-]+)+(?:/.*)?$"     # host names, dotted codes
    r"|[^@\s]+@[^@\s]+\.[A-Za-z]{2,}$"                   # email addresses
    r"|[A-Z][a-z0-9]+(?:[A-Z][A-Za-z0-9]*)+$"            # ApiBase, OutDir, AppKey
    r"|[A-Za-z]+(?:-[A-Za-z0-9]+){2,}$"                  # claude-haiku-4-5-...
    r"|SIM[A-Z]*-[A-Za-z]+$"                             # SIMF-Prod, SIM-RNSF
    r")")

# Values the pattern above would claim but which are ordinary numbers: "0.5"
# is not a dotted code, and rendering it in a different face from the "20" and
# "5" beside it looks like a mistake.
NOT_TECHNICAL = re.compile(r"^[0-9]+(?:\.[0-9]+)?$")


def is_technical(text):
    """Whether a value is machine text that must render left to right."""
    value = text.strip()
    if not value or NOT_TECHNICAL.match(value):
        return False
    return (value.startswith(TECHNICAL_PREFIXES)
            or value.endswith(TECHNICAL_SUFFIXES)
            or bool(TECHNICAL_PATTERN.match(value)))


class Builder:
    def __init__(self, lang, book, pages, redirected=frozenset()):
        self.lang = lang
        self.rtl = lang == "ar"
        self.book = book
        self.pages = pages
        # Slugs the capture proved unreachable: the browser was sent there and
        # landed somewhere else. They get an explanation, not a picture of
        # whatever page answered instead.
        self.redirected = redirected
        # The resource strings in THIS volume's language, so a grid column can
        # be named with the word the page itself puts in the header.
        self.strings = read_resx(lang)
        self.figure_number = 0
        self.missing = []
        self.document = Document()
        style_rtl(self.document, self.rtl)
        self._setup_section()

    # ------------------------------------------------------------ helpers --

    def _t(self, value):
        """Pick this volume's half of an {en, ar} pair."""
        if isinstance(value, dict):
            if self.lang not in value:
                raise KeyError(f"block is missing the '{self.lang}' text: {value}")
            return value[self.lang]
        return value

    def _setup_section(self):
        section = self.document.sections[0]
        section_rtl(section, self.rtl)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_rtl(footer, self.rtl)
        add_field(footer, "PAGE")

    def para(self, text, style=None, size=10.5, bold=False, color=None,
             mono=False, align=None, space_after=6):
        paragraph = self.document.add_paragraph(style=style)
        run = paragraph.add_run(text)
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color
        run_rtl(run, self.rtl)
        run_font(run,
                 latin=MONO_FONT if mono else LATIN_FONT,
                 complex_script=ARABIC_FONT,
                 size_pt=size)
        paragraph_rtl(paragraph, self.rtl)
        if align is not None:
            paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(space_after)
        return paragraph

    def mixed(self, pieces, style=None, size=10.5, color=None,
              align=None, space_after=6):
        """A paragraph whose Latin pieces keep their own direction.

        `pieces` is a list of (text, is_machine_text). Each becomes its own run,
        and a machine-text run has right-to-left cleared so a route or a path
        renders as written inside Arabic prose.
        """
        paragraph = self.document.add_paragraph(style=style)
        for text, machine in pieces:
            if not text:
                continue
            run = paragraph.add_run(text)
            if color is not None:
                run.font.color.rgb = color
            run_rtl(run, self.rtl and not machine)
            run_font(run,
                     latin=MONO_FONT if machine else LATIN_FONT,
                     complex_script=ARABIC_FONT,
                     size_pt=size)
        paragraph_rtl(paragraph, self.rtl)
        if align is not None:
            paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(space_after)
        return paragraph

    # ------------------------------------------------------------- blocks --

    def heading(self, text, level):
        paragraph = self.document.add_paragraph(style=f"Heading {level}")
        run = paragraph.add_run(text)
        run_rtl(run, self.rtl)
        run_font(run, size_pt={1: 18, 2: 14, 3: 12, 4: 11}.get(level, 11))
        paragraph_rtl(paragraph, self.rtl)
        paragraph.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
        paragraph.paragraph_format.space_after = Pt(6)

    def bullets(self, items):
        for item in items:
            self.para(item, style="List Bullet", size=10.5, space_after=3)

    def table(self, headers, rows, widths=None):
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table_rtl(table, self.rtl)
        table.autofit = True

        for cell, heading in zip(table.rows[0].cells, headers):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(str(heading))
            run.bold = True
            run.font.color.rgb = ACCENT
            run_rtl(run, self.rtl)
            run_font(run, size_pt=9)
            paragraph_rtl(paragraph, self.rtl)

        for row in rows:
            if len(row) != len(headers):
                raise ValueError(
                    f"table row has {len(row)} cells but the header has "
                    f"{len(headers)}: {row}")
            cells = table.add_row().cells
            for cell, value in zip(cells, row):
                cell.text = ""
                paragraph = cell.paragraphs[0]
                text = "" if value is None else str(value)
                # A route, a path or a permission code is Latin even in the
                # Arabic volume; rendering it right-to-left moves its slashes.
                technical = is_technical(text)
                run = paragraph.add_run(text)
                run_rtl(run, self.rtl and not technical)
                run_font(run,
                         latin=MONO_FONT if technical else LATIN_FONT,
                         size_pt=8.5)
                paragraph_rtl(paragraph, self.rtl and not technical)
                paragraph.paragraph_format.space_after = Pt(2)
        self.para("", size=4, space_after=2)
        return table

    def figure(self, image_name, caption, machine_suffix=""):
        # The Arabic volume must show the Arabic interface. Authored chapters name
        # an image without a language suffix, so prefer the "-ar" capture of the
        # same screen when one exists and fall back to the English one when it
        # does not - the three one-shot sign-in screens were captured once, in
        # English, and cannot be recaptured.
        path = None
        if self.rtl and not image_name.endswith("-ar"):
            arabic = SHOTS / f"{image_name}-ar.png"
            if arabic.exists():
                path = arabic
            elif image_name not in ENGLISH_ONLY:
                # Falling back to the English screen for anything else would put
                # an English interface in the Arabic book and still report a
                # clean build, which is how a partly-failed Arabic capture would
                # go unnoticed.
                self.missing.append(image_name + "-ar")
                return
        if path is None:
            path = SHOTS / f"{image_name}.png"
        if not path.exists():
            self.missing.append(image_name)
            return
        self.figure_number += 1
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run().add_picture(str(path), width=Inches(6.3))

        prefix = ("شكل" if self.rtl else "Figure") + f" {self.figure_number} — "
        self.mixed([(prefix + caption, False), (machine_suffix, True)],
                   style="Caption", size=8.5, color=MUTED,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    def note(self, text):
        self.para(text, size=10, color=MUTED, space_after=8)

    def page_break(self):
        paragraph = self.document.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    # -------------------------------------------------------------- front --

    def cover(self):
        meta = self.book["meta"]
        for _ in range(4):
            self.para("", size=12)
        self.para(self._t(meta["title"]), style="Title", size=26,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        self.para(self._t(meta["subtitle"]), size=13, color=MUTED,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)
        rows = [[self._t(k), self._t(v)] for k, v in meta["facts"]]
        self.table([self._t(meta["factHeaders"][0]), self._t(meta["factHeaders"][1])], rows)
        self.page_break()

        self.heading(self._t(meta["contentsTitle"]), 1)
        toc = self.document.add_paragraph()
        paragraph_rtl(toc, self.rtl)
        add_field(toc, r'TOC \o "1-2" \h \z \u')
        self.page_break()

    # ------------------------------------------------------------- render --

    def render(self):
        self.cover()
        for chapter in self.book["chapters"]:
            self.heading(self._t(chapter["title"]), 1)
            for block in chapter["blocks"]:
                self.block(block)
            self.page_break()
        self.reference_section()
        mark_dirty_fields(self.document)
        return self.document

    def block(self, block):
        kind = block["t"]
        if kind in ("h2", "h3", "h4"):
            self.heading(self._t(block), int(kind[1]))
        elif kind == "p":
            self.para(self._t(block))
        elif kind == "note":
            self.note(self._t(block))
        elif kind == "code":
            self.para(self._t(block), size=9, mono=True)
        elif kind == "bullets":
            self.bullets(self._t(block))
        elif kind == "table":
            self.table(self._t(block["headers"]), self._t(block["rows"]))
        elif kind == "figure":
            self.figure(block["image"], self._t(block["caption"]))
        elif kind == "pagebreak":
            self.page_break()
        else:
            raise ValueError(f"unknown block type: {kind}")

    # --------------------------------------------------- generated part ---

    def reference_section(self):
        """One entry per Control Panel page, in navigation order.

        Generated from page-model.json rather than authored, so the route, the
        implementing file, the permission and the menu label in both languages
        are whatever the code says today.
        """
        strings = self.book["reference"]
        self.heading(self._t(strings["title"]), 1)
        self.para(self._t(strings["intro"]))

        headers = [self._t(h) for h in strings["headers"]]
        current_group = None
        in_nav = [p for p in self.pages if p.get("inNavigation")]

        for page in in_nav:
            group = page.get("navGroupKey")
            if group != current_group:
                current_group = group
                label = page.get("navGroupAr") if self.rtl else page.get("navGroupEn")
                self.heading(label or group, 2)
            self.page_entry(page, headers, strings)

        self.page_break()
        self.heading(self._t(strings["nonNavTitle"]), 2)
        self.para(self._t(strings["nonNavIntro"]))
        for page in [p for p in self.pages if not p.get("inNavigation")]:
            self.page_entry(page, headers, strings)

    def page_depth(self, page, strings):
        """What the page offers, read out of the page itself.

        Every row here is extracted from the razor and its code-behind, so it
        describes the page as it is now. Nothing is written by hand, which is
        why it can be given for all 114 pages instead of the dozen somebody
        would realistically keep up to date.
        """
        actions = page.get("actions") or []
        gate_list = page.get("actionGates") or []

        # Which grid permission covers which toolbar action. Add, Duplicate and
        # Paste share one, because they are the same verb: all three create a
        # record. Only these have a grid parameter; anything else the page wraps
        # itself, and those are listed separately below.
        covered_by = {
            "Add": "Add, Duplicate and Paste",
            "Duplicate": "Add, Duplicate and Paste",
            "Paste": "Add, Duplicate and Paste",
            "Edit": "Edit",
            "Delete": "Delete",
            "Delete selected": "Delete",
            "Approve selected": "Approve",
            "Reject selected": "Reject",
            "Import": "Import",
            "Export": "Export",
        }
        grid_gates = {g["covers"]: g["permission"] for g in gate_list
                      if g["covers"] != "An action on the page"}
        # Kept as a LIST: a page may wrap several different buttons, each with
        # its own permission, and collapsing them into a dictionary keyed by the
        # label would report only the last one.
        page_gates = [g["permission"] for g in gate_list
                      if g["covers"] == "An action on the page"]

        if actions or page_gates:
            self.para(self._t(strings["actionsTitle"]), bold=True, size=10, space_after=3)
            rows = []
            for action in actions:
                name = action["name"]
                shown = self.label_for(action.get("labelKey") or "") or name
                # Two callbacks can share one label - Delete and Delete selected
                # both say "Delete" - so the row keeps the distinction the
                # toolbar makes by position, not by wording alone.
                if name.endswith(" selected") and not shown.lower().endswith("selected"):
                    shown = f"{shown} ({self._t(strings['bulk'])})"
                rows.append([shown,
                             grid_gates.get(covered_by.get(name, ""))
                             or self._t(strings["ungated"])])
            for permission in sorted(set(page_gates)):
                rows.append([self._t(strings["pageButton"]), permission])
            self.table([self._t(h) for h in strings["actionsHeaders"]], rows)

        columns = page.get("columns") or []
        if columns:
            names = []
            for column in columns:
                key = column.get("headerKey") or ""
                names.append(self.label_for(key) or column.get("key", ""))
            self.para(f"{self._t(strings['columnsLabel'])}: " + "  ·  ".join(names),
                      size=9.5, color=MUTED, space_after=4)

        calls = page.get("calls") or []
        if calls:
            self.mixed([(self._t(strings["callsLabel"]) + ": ", False),
                        ("  ".join(calls), True)],
                       size=8.5, color=MUTED, space_after=8)

    def label_for(self, key):
        """A resource string in this volume's language, or None."""
        return self.strings.get(key)

    def page_entry(self, page, headers, strings):
        label = (page.get("labelAr") if self.rtl else page.get("labelEn")) \
            or (page.get("titleAr") if self.rtl else page.get("titleEn")) \
            or page["route"]
        self.mixed([(label + "  (", False), (page["route"], True), (")", False)],
                   style="Heading 3", size=12)

        access = page.get("access")
        gate = page.get("permission") or (
            self._t(ACCESS_TEXT[access]) if access in ACCESS_TEXT else "-")

        rows = [
            [self._t(strings["rowRoute"]), page["route"]],
            [self._t(strings["rowFile"]), page.get("razor") or "-"],
            [self._t(strings["rowCodeBehind"]), page.get("codeBehind") or "-"],
            [self._t(strings["rowPermission"]), gate],
            [self._t(strings["rowNavPermission"]), page.get("navPermission") or "-"],
        ]
        if page.get("isStub"):
            rows.append([self._t(strings["rowStub"]), self._t(strings["stubYes"])])
        self.table(headers, rows)

        # A page reached from inside another page has no menu label, and the
        # heading already carries its address - repeating it in the caption
        # reads as a mistake rather than as emphasis.
        caption_label = page.get("labelAr") if self.rtl else page.get("labelEn")
        self.page_depth(page, strings)

        if page["slug"] in self.redirected:
            self.note(self._t(strings["redirected"]))
            return
        self.figure(f"cp-{page['slug']}-default",
                    f"{caption_label} — " if caption_label else "",
                    machine_suffix=page["route"])


def read_resx(lang):
    """name -> value for one language's Control Panel strings."""
    import xml.etree.ElementTree as ET
    name = "Strings.resx" if lang == "en" else f"Strings.{lang}.resx"
    path = REPO / "src/ControlPanel/SIMF.ControlPanel/Resources" / name
    values = {}
    for data in ET.parse(path).getroot().findall("data"):
        key, value = data.get("name"), data.find("value")
        if key and value is not None and value.text is not None:
            values[key] = value.text
    return values


def redirected_slugs(lang):
    """Slugs the capture recorded as unreachable for a signed-in reader.

    Five routes redirect: the two account-state pages go to the dashboard and
    the three mid-sign-in pages go back to the sign-in form. Photographing them
    produced a picture of the wrong page filed under the right name, and the
    build passed because the FILE existed.
    """
    report = SHOTS / f"capture-report-sweep-{lang}.json"
    if not report.exists():
        return frozenset()
    entries = json.loads(report.read_text(encoding="utf-8"))
    return frozenset(e["slug"] for e in entries if e.get("redirected"))


def main():
    book = json.loads((SOURCE / "book.json").read_text(encoding="utf-8"))
    pages = json.loads((SOURCE / "page-model.json").read_text(encoding="utf-8"))

    OUT.mkdir(parents=True, exist_ok=True)
    problems, rendered = [], []
    for lang, filename in (("en", "SIMF-CP-Operations-Manual-EN.docx"),
                           ("ar", "SIMF-CP-Operations-Manual-AR.docx")):
        builder = Builder(lang, book, pages, redirected_slugs(lang))
        document = builder.render()
        rendered.append((lang, OUT / filename, document, builder))
        if builder.missing:
            problems.append((lang, builder.missing))

    # Save only when BOTH volumes are complete. Saving as each one renders would
    # overwrite a good published manual with a holed one and then exit non-zero.
    if not problems:
        for lang, path, document, builder in rendered:
            document.save(path)
            print(f"{lang}: {path.relative_to(REPO)}  "
                  f"({path.stat().st_size // 1024} KB, {builder.figure_number} figures)")

    if problems:
        print("\nMISSING SCREENSHOTS - the build is not clean:")
        for lang, names in problems:
            unique = sorted(set(names))
            print(f"  {lang}: {len(unique)} missing")
            for name in unique[:25]:
                print(f"    {name}.png")
            if len(unique) > 25:
                print(f"    ... and {len(unique) - 25} more")
        raise SystemExit(1)
    print("\nEvery referenced screenshot exists.")


if __name__ == "__main__":
    main()
