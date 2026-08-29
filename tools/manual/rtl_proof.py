"""Prove the RTL Word output before the manual is built on top of it.

Generates one English (LTR) and one Arabic (RTL) two-page document exercising
exactly the features the manual needs and that python-docx does not support
natively: section direction, paragraph direction, run direction, mirrored table
columns, the complex-script font, an image with a caption, a TOC field and a
page-number footer.

Run:  python tools/manual/rtl_proof.py
"""

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx_kit import (  # noqa: E402
    add_field, mark_dirty_fields, paragraph_rtl, run_font, run_rtl,
    section_rtl, style_rtl, table_rtl,
)

OUT = Path(__file__).resolve().parents[2] / ".tmp" / "manual-proof"

EN = {
    "title": "SIMF Control Panel — RTL proof",
    "h1": "Creating an administrator",
    "body": "Open Access control, then Admins, and choose Add. The form asks "
            "for an email address, a display name and the roles the new "
            "administrator should hold.",
    "table_head": ["Field", "Required", "Maximum length"],
    "table_rows": [
        ["Email address", "Yes", "256"],
        ["Display name", "Yes", "2 to 128"],
        ["Roles", "No", "—"],
    ],
    "caption": "Figure 1 — the Add user form on /admin/admins",
}

# Every Arabic string below is a UI label taken from the Control Panel's own
# Strings.ar.resx, not a hand translation: Module.AdminAdmins, Nav.AccessControl,
# Admin.CreateUser.Email, Admin.CreateUser.DisplayName, Admin.CreateUser.RolesLabel.
AR = {
    "title": "لوحة تحكم SIMF — إثبات الاتجاه من اليمين إلى اليسار",
    "h1": "المسؤولون",
    "body": "افتح التحكم في الوصول، ثم المسؤولون، واختر إضافة. يطلب النموذج "
            "البريد الإلكتروني والاسم المعروض والأدوار التي يحملها المسؤول الجديد.",
    "table_head": ["الحقل", "مطلوب", "الحد الأقصى"],
    "table_rows": [
        ["البريد الإلكتروني", "نعم", "256"],
        ["الاسم المعروض", "نعم", "2 إلى 128"],
        ["الأدوار", "لا", "—"],
    ],
    "caption": "شكل 1 — نموذج إضافة مستخدم في /admin/admins",
}


def para(document, text, rtl, style=None, size=None, bold=False):
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.bold = bold
    run_rtl(run, rtl)
    run_font(run, size_pt=size)
    paragraph_rtl(paragraph, rtl)
    return paragraph


def build(strings, rtl, image, path):
    document = Document()
    style_rtl(document, rtl)
    section = document.sections[0]
    section_rtl(section, rtl)

    para(document, strings["title"], rtl, style="Title")
    para(document, "", rtl)

    # Table of contents — a field Word evaluates on open.
    toc = document.add_paragraph()
    paragraph_rtl(toc, rtl)
    add_field(toc, r'TOC \o "1-3" \h \z \u')
    document.add_page_break()

    para(document, strings["h1"], rtl, style="Heading 1")
    para(document, strings["body"], rtl, size=11)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table_rtl(table, rtl)
    for cell, heading in zip(table.rows[0].cells, strings["table_head"]):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(heading)
        run.bold = True
        run_rtl(run, rtl)
        run_font(run, size_pt=10)
        paragraph_rtl(paragraph, rtl)
    for row in strings["table_rows"]:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(value)
            run_rtl(run, rtl)
            run_font(run, size_pt=10)
            paragraph_rtl(paragraph, rtl)

    para(document, "", rtl)
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(image), width=Inches(5.5))
    para(document, strings["caption"], rtl, style="Caption", size=9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_rtl(footer, rtl)
    add_field(footer, "PAGE")

    mark_dirty_fields(document)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def placeholder_image(path):
    """A neutral placeholder so the proof does not depend on a capture run."""
    from PIL import Image, ImageDraw
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1200, 500), "#eef1f5")
    draw = ImageDraw.Draw(image)
    draw.rectangle([4, 4, 1195, 495], outline="#5b6b7f", width=3)
    draw.text((40, 40), "screenshot placeholder", fill="#31415a")
    image.save(path)
    return path


if __name__ == "__main__":
    image = placeholder_image(OUT / "placeholder.png")
    for strings, rtl, name in ((EN, False, "proof-EN.docx"), (AR, True, "proof-AR.docx")):
        written = build(strings, rtl, image, OUT / name)
        print(f"wrote {written} ({written.stat().st_size} bytes)")
